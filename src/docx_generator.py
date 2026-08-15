"""Renders the same structured resume data as pdf_generator, but to DOCX.

Section order and content mirror the PDF templates exactly. Rather than
restating the rules for which fields appear (school vs college, whether a
CGPA/grade toggle is on, which project rows have content), this module
imports the same helpers pdf_generator uses, so the two formats cannot
drift apart.

Word's own list styles do the bullets; nothing here tries to reproduce the
PDF's manual glyph drawing.
"""

from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

from src.pdf_generator import (
    _clean_items,
    _contact_line,
    _education_lines,
    _education_with_content,
    _links_line,
    _project_bullets,
    _projects_with_content,
    _work_entry_bullets,
)

CLASSIC = "classic"
PROFESSIONAL = "professional"

SERIF_FONT = "Georgia"          # Professional body text
NAME_SIZE_PT = 20
HEADING_SIZE_PT = 12
BODY_SIZE_PT = 10.5
CONTACT_SIZE_PT = 10


def _normalise_template(template):
    name = str(template or CLASSIC).strip().lower()
    return PROFESSIONAL if name.startswith("pro") else CLASSIC


def _text(value):
    return str(value or "").strip()


def _set_bottom_border(paragraph):
    """Thin rule beneath a paragraph, mirroring the ruled PDF headings."""
    p_pr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")        # eighths of a point
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    borders.append(bottom)
    p_pr.append(borders)


def _add_right_tab(paragraph, section):
    """Right-aligned tab stop at the right margin."""
    width = section.page_width - section.left_margin - section.right_margin
    paragraph.paragraph_format.tab_stops.add_tab_stop(
        width, WD_TAB_ALIGNMENT.RIGHT
    )


def _para(doc, text="", size=BODY_SIZE_PT, bold=False, italic=False,
          align=None, space_after=2, style=None):
    paragraph = doc.add_paragraph(style=style)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(space_after)
    if align is not None:
        paragraph.alignment = align
    if text:
        run = paragraph.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.size = Pt(size)
    return paragraph


def _heading(doc, title, professional):
    paragraph = _para(doc, title.upper(), size=HEADING_SIZE_PT, bold=True,
                      space_after=2)
    paragraph.paragraph_format.space_before = Pt(10)
    _set_bottom_border(paragraph)
    return paragraph


def _bullet(doc, text, size=BODY_SIZE_PT):
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(1)
    run = paragraph.add_run(_text(text))
    run.font.size = Pt(size)
    return paragraph


def _entry_header(doc, section, primary, dates, size=BODY_SIZE_PT):
    """Bold primary on the left with dates right-aligned on the same line."""
    primary, dates = _text(primary), _text(dates)
    if not primary and not dates:
        return None
    paragraph = _para(doc, space_after=1)
    if dates:
        _add_right_tab(paragraph, section)
    if primary:
        run = paragraph.add_run(primary)
        run.bold = True
        run.font.size = Pt(size)
    if dates:
        run = paragraph.add_run("\t" + dates)
        run.font.size = Pt(size - 0.5)
    return paragraph


def _apply_font(doc, font_name):
    """Force a font across body text, styles included.

    Setting only the Normal style leaves List Bullet (and any run that
    carries its own rFonts) on the theme font, so walk the runs too.
    """
    for style_name in ("Normal", "List Bullet"):
        try:
            style = doc.styles[style_name]
        except KeyError:
            continue
        style.font.name = font_name
        r_pr = style.element.get_or_add_rPr()
        r_fonts = r_pr.get_or_add_rFonts()
        for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
            r_fonts.set(qn(attr), font_name)

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.name = font_name
            r_pr = run._element.get_or_add_rPr()
            r_fonts = r_pr.get_or_add_rFonts()
            for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
                r_fonts.set(qn(attr), font_name)


def generate_resume_docx(personal_info, work_history, education, skills,
                         achievements, certifications, projects, links,
                         template=CLASSIC):
    """Build the resume as a .docx and return it as bytes.

    ``links`` is accepted separately for convenience; when it is falsy the
    links already carried on ``personal_info`` are used instead, so callers
    can pass either shape.

    ``certifications`` is rendered whenever it holds a non-empty entry - the
    caller applies its own "include certifications" toggle by passing an
    empty list, exactly as the PDF flow does.
    """
    professional = _normalise_template(template) == PROFESSIONAL

    doc = Document()
    section = doc.sections[0]

    header_info = dict(personal_info or {})
    if links:
        header_info["links"] = links

    align = WD_ALIGN_PARAGRAPH.CENTER if professional else WD_ALIGN_PARAGRAPH.LEFT

    # ----- Name & contact -----

    name = _text(header_info.get("name")) or "Resume"
    _para(doc, name, size=NAME_SIZE_PT, bold=True, align=align, space_after=2)

    contact_line = _contact_line(header_info)
    if contact_line:
        _para(doc, contact_line, size=CONTACT_SIZE_PT, align=align, space_after=1)

    links_line = _links_line(header_info)
    if links_line:
        _para(doc, links_line, size=CONTACT_SIZE_PT, align=align, space_after=1)

    # ----- Experience -----

    work_entries = [
        entry for entry in (work_history or [])
        if isinstance(entry, dict)
        and (entry.get("company") or entry.get("role") or _work_entry_bullets(entry))
    ]
    if work_entries:
        _heading(doc, "Experience", professional)
        for entry in work_entries:
            if professional:
                _entry_header(doc, section, entry.get("company"), entry.get("dates"))
                role = _text(entry.get("role"))
                if role:
                    _para(doc, role, italic=True, space_after=1)
            else:
                title = " - ".join(
                    part for part in (_text(entry.get("role")),
                                      _text(entry.get("company"))) if part
                )
                if title:
                    _para(doc, title, bold=True, space_after=1)
                dates = _text(entry.get("dates"))
                if dates:
                    _para(doc, dates, size=BODY_SIZE_PT - 0.5, italic=True,
                          space_after=1)
            for bullet in _work_entry_bullets(entry):
                _bullet(doc, bullet)

    # ----- Education (school vs college handled by the shared helper) -----

    education_entries = _education_with_content(education)
    if education_entries:
        _heading(doc, "Education", professional)
        for entry in education_entries:
            edu_name, secondary, dates = _education_lines(entry)
            if professional:
                _entry_header(doc, section, edu_name, dates)
                if secondary:
                    _para(doc, secondary, italic=True, space_after=1)
            else:
                if edu_name:
                    _para(doc, edu_name, bold=True, space_after=1)
                if secondary:
                    _para(doc, secondary, space_after=1)
                if dates:
                    _para(doc, dates, size=BODY_SIZE_PT - 0.5, italic=True,
                          space_after=1)

    # ----- Skills -----

    skill_items = _clean_items(skills)
    if skill_items:
        _heading(doc, "Skills", professional)
        _para(doc, ", ".join(skill_items), space_after=2)

    # ----- Projects -----

    project_entries = _projects_with_content(projects)
    if project_entries:
        _heading(doc, "Projects", professional)
        for project in project_entries:
            project_name = _text(project.get("name"))
            if professional:
                _entry_header(doc, section, project_name, "")
            elif project_name:
                _para(doc, project_name, bold=True, space_after=1)
            tech_stack = _text(project.get("tech_stack"))
            if tech_stack:
                _para(doc, tech_stack, size=BODY_SIZE_PT - 0.5, italic=True,
                      space_after=1)
            for bullet in _project_bullets(project):
                _bullet(doc, bullet)
            link = _text(project.get("link"))
            if link:
                _para(doc, f"Link: {link}", size=BODY_SIZE_PT - 0.5, space_after=1)

    # ----- Certifications (caller passes an empty list when toggled off) -----

    certification_items = _clean_items(certifications)
    if certification_items:
        _heading(doc, "Certifications", professional)
        for certification in certification_items:
            _bullet(doc, certification)

    # ----- Achievements -----

    achievement_items = _clean_items(achievements)
    if achievement_items:
        _heading(doc, "Achievements", professional)
        for achievement in achievement_items:
            _bullet(doc, achievement)

    if professional:
        _apply_font(doc, SERIF_FONT)

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
